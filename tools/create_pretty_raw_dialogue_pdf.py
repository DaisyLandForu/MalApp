from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from create_raw_dialogue_pdf import MESSAGES


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "弱标签与引擎评分机制_对话原文美化版.docx"


def set_run_font(run, size: float = 10.0, color: str | None = None, bold: bool = False, font: str = "Microsoft YaHei") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 160, right: int = 160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))


def looks_like_code(block: str) -> bool:
    stripped = block.strip()
    if stripped.startswith("```") or stripped.startswith(("{", "[", "def ", "if ", "elif ", "else:", "GROUP BY", "HAVING")):
        return True
    markers = ("->", "=>", " = ", ">= ", "<= ", "python ", ".\\", "```text", "```json", "```powershell")
    return any(marker in stripped for marker in markers)


def add_code_block(cell, text: str) -> None:
    table = cell.add_table(rows=1, cols=1)
    set_table_width(table, 8200)
    code_cell = table.cell(0, 0)
    set_cell_shading(code_cell, "F5F7FA")
    set_cell_border(code_cell, "DDE5EE")
    set_cell_margins(code_cell, 90, 90, 140, 140)
    p = code_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, 8.3, "1F4D78", font="Consolas")


def add_text_block(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, 9.8, "1F2937")


def add_message_body(cell, content: str) -> None:
    cell.text = ""
    for block in content.split("\n\n"):
        if not block.strip():
            continue
        if looks_like_code(block):
            add_code_block(cell, block)
        else:
            add_text_block(cell, block)


def add_message(doc: Document, index: int, role: str, content: str) -> None:
    is_user = role == "用户"
    role_fill = "E7F0FF" if is_user else "EAF7EF"
    body_fill = "FBFDFF" if is_user else "FCFFFD"
    border = "C8D9F2" if is_user else "CAE6D3"
    accent = "1D4ED8" if is_user else "137333"

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    table.columns[0].width = Inches(0.78)
    table.columns[1].width = Inches(5.72)
    role_cell, body_cell = table.rows[0].cells
    role_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    body_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for cell in (role_cell, body_cell):
        set_cell_border(cell, border)
        set_cell_margins(cell)
    set_cell_shading(role_cell, role_fill)
    set_cell_shading(body_cell, body_fill)

    p = role_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(role)
    set_run_font(r, 9.2, accent, True)
    p2 = role_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = p2.add_run(f"#{index}")
    set_run_font(n, 8.0, "6B7280")

    add_message_body(body_cell, content)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.8)
    normal.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("弱标签与引擎评分机制 · 对话原文"), 8.5, "6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("原文打印版 · 未整理改写"), 8.5, "9CA3AF")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    set_run_font(title.add_run("弱标签与引擎评分机制"), 20, "0B2545", True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("对话原文美化打印版"), 10.5, "6B7280")

    for index, (role, content) in enumerate(MESSAGES, start=1):
        add_message(doc, index, role, content)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
