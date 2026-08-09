from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "outputs" / "resume" / "王雪-v4_算法工程师.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, bold: bool = False, size: float = 9.5, color: str | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc: Document, title: str, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.08
    add_run(p, title, bold=True)
    add_run(p, text)


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E78")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    add_run(p, title, bold=True, size=12, color="1F4E78")


def add_project(doc: Document, name: str, role: str, date: str, bullets: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    cells = table.rows[0].cells
    set_cell_width(cells[0], 7200)
    set_cell_width(cells[1], 2160)
    p = cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_run(p, name, bold=True, size=10.5)
    p2 = cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p2, date, size=9, color="5B6573")
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(1)
    add_run(p3, role, bold=True, size=9, color="5B6573")
    for title, body in bullets:
        add_bullet(doc, title, body)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "王雪", bold=True, size=22, color="17365D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "算法工程师 ｜ AI 应用开发", bold=True, size=11, color="1F4E78")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "北京 · 海淀区  |  15064265610  |  zy2402425@buaa.edu.cn  |  硕士在读", size=9, color="4A5568")

    add_section(doc, "个人概述")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    add_run(p, "具备机器学习、信号处理与 AI 应用工程背景；能够完成从多源数据接入、特征预处理、模型推理与评估，到本地服务和可视化工作台的完整闭环。近期聚焦大模型应用、多智能体协同与恶意 APP 风险研判。", size=9.5)

    add_section(doc, "教育经历")
    edu = doc.add_table(rows=2, cols=2)
    edu.autofit = False
    edu.allow_autofit = False
    for row in edu.rows:
        set_cell_width(row.cells[0], 7200)
        set_cell_width(row.cells[1], 2160)
    p = edu.rows[0].cells[0].paragraphs[0]
    add_run(p, "北京航空航天大学  |  电子科学与技术  |  硕士", bold=True, size=10)
    p = edu.rows[0].cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "2024/09 - 至今", size=9, color="5B6573")
    p = edu.rows[1].cells[0].paragraphs[0]
    add_run(p, "山东大学  |  通信工程  |  本科（GPA 3.88，专业前 5%）", bold=True, size=10)
    p = edu.rows[1].cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "2020/09 - 2024/06", size=9, color="5B6573")

    add_section(doc, "专业技能")
    skills = [
        ("AI / 算法", "Python、PyTorch、Transformers、机器学习、深度学习、信号处理与模型评估。"),
        ("大模型应用", "Qwen2.5 本地推理、提示词设计、结构化 JSON 输出、双模型多轮辩论与多智能体协同。"),
        ("数据工程", "JSON/XML/Excel 解析、字段归一化、IOC 提取、SQLite、弱标注 train/val/test 构建与参数网格搜索。"),
        ("工程化", "本地 HTTP API、任务优先级队列、结果缓存、可追溯研判报告、Web 工作台。"),
    ]
    for title, body in skills:
        add_bullet(doc, f"{title}：", body)

    add_section(doc, "项目经历")
    add_project(
        doc,
        "基于双模型辩论与多智能体协同的恶意 APP 研判引擎",
        "AI 应用算法与数据工程开发",
        "2026/05 - 至今",
        [
            ("研判架构：", "构建恶意 APP 本地研判 MVP，融合 360/cm 引擎结果、APP 标签、静态特征和通联 IOC，输出风险分数、恶意/可疑/良性结论及可追溯证据链。"),
            ("多智能体与大模型：", "实现静态分析、情报溯源、仿冒研判、业务打标 4 个证据智能体；基于本地 Qwen2.5 完成模型甲/乙初判、互相质询和反驳，生成 Engine C 仲裁结果。"),
            ("数据管道：", "接入 360/cm、APP_md5、人工标注冲突样本、APP 主画像和通联地址等多源数据；完成 186 个字段注册及 5 万+特征记录的归一化、静态特征包和网络 IOC 包构建。"),
            ("策略评估：", "实现 WEC 加权融合、阈值/权重网格评估、高分歧样本自动拉取、人工复核优先级队列、SQLite 缓存及 Web 工作台展示。"),
        ],
    )
    add_project(
        doc,
        "基于 U-Sleep 的 NT1 智能筛查系统",
        "算法开发与模型实现",
        "2025/12",
        [
            ("模型优化：", "设计 4 通道多分辨率睡眠分期方案，结合小波与注意力机制，分期准确率由 71% 提升至 75%。"),
            ("筛查评估：", "面向低发病率罕见病筛查优化判别策略，在 99.0% 特异性下实现 96.3% 灵敏度；完成预处理、分期、特征提取、分类和不确定性评估链路。"),
        ],
    )
    add_project(
        doc,
        "智能优化算法与联合仿真自动化研究",
        "算法开发与流程实现",
        "2024/12 - 2025/09",
        [
            ("优化与仿真：", "基于 Python 完成多种智能优化算法的对比分析和参数寻优；使用 MATLAB-CST 构建参数驱动、自动求解与结果分析的联合仿真流程。"),
        ],
    )

    add_section(doc, "科研经历")
    add_bullet(doc, "AI 成像与反演：", "围绕 AI 成像、深度学习算法、离散损失优化、自动微分、有限元建模与物理约束反演开展研究。")
    add_bullet(doc, "论文成果：", "参与多篇国际会议与 SCI 期刊论文工作；IEEE TIM（SCI Q1）投稿 1 篇，ACES-China 2025 与 IEEE NEMO 2025 论文录用。")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "王雪 · 算法工程师 / AI 应用开发", size=8, color="7A8493")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
